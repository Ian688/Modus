"""Office document tools (Phase A3): Excel/Word/PPT for everyday users.

These parse binary formats (xlsx/docx/pptx) that the text tools cannot read.
Read tools auto-ALLOW and disclose workspace content; write tools are
approval-gated.  Multi-thousand-row Excel works (streaming read_only + iter_rows).
"""

from __future__ import annotations

import pytest

from modus.config import ModusConfig
from modus.policy.approval import ApprovalDecision, ApprovalPolicy
from modus.tools.base import ToolContext
from modus.tools.builtins import get_builtin_tools
from modus.tools.office import (
    excel_analyze,
    excel_query,
    pptx_build,
    pptx_extract,
    word_edit,
    word_extract,
)


@pytest.fixture
def office_ws(tmp_path, monkeypatch):
    from pathlib import Path

    monkeypatch.setattr(Path, "home", classmethod(lambda cls: tmp_path))
    ws = tmp_path / "ws"
    ws.mkdir()
    return ws


def _ctx(ws):
    return ToolContext(cwd=str(ws), workspace_root=str(ws), config=ModusConfig())


def _make_excel(ws, rows=5000):
    import openpyxl

    wb = openpyxl.Workbook()
    sh = wb.active
    sh.append(["name", "score", "city"])
    for i in range(rows):
        sh.append([f"user{i}", i % 100, "beijing" if i % 2 else "shanghai"])
    path = ws / "data.xlsx"
    wb.save(path)
    return path


# ── declaration contract ──


def test_office_tools_declared():
    tools = {t.name: t for t in get_builtin_tools()}
    for name in ("excel_analyze", "excel_query", "word_extract", "word_edit",
                 "pptx_extract", "pptx_build"):
        assert name in tools


def test_read_tools_auto_allow_and_disclose_content():
    tools = {t.name: t for t in get_builtin_tools()}
    policy = ApprovalPolicy(ModusConfig().policy)
    for name in ("excel_analyze", "excel_query", "word_extract", "pptx_extract"):
        tool = tools[name]
        assert tool.is_read_only is True
        assert tool.data_disclosure == "workspace_content"
        assert policy.evaluate(tool) is ApprovalDecision.ALLOW
        assert "filesystem" in tool.capabilities


def test_write_tools_approval_gated():
    tools = {t.name: t for t in get_builtin_tools()}
    policy = ApprovalPolicy(ModusConfig().policy)
    for name in ("word_edit", "pptx_build"):
        tool = tools[name]
        assert tool.is_read_only is False
        assert tool.danger_level == "medium"
        assert policy.evaluate(tool) is ApprovalDecision.ASK


# ── Excel ──


@pytest.mark.asyncio
async def test_excel_analyze_large_sheet(office_ws):
    _make_excel(office_ws)
    result = await excel_analyze({"path": "data.xlsx"}, _ctx(office_ws))
    assert not result.is_error
    assert "5001 rows" in result.content
    assert "Header: ['name', 'score', 'city']" in result.content
    assert "user0" in result.content


@pytest.mark.asyncio
async def test_excel_analyze_numeric_stats(office_ws):
    _make_excel(office_ws)
    result = await excel_analyze({"path": "data.xlsx"}, _ctx(office_ws))
    assert "Numeric stats:" in result.content
    assert "score" in result.content


@pytest.mark.asyncio
async def test_excel_query_gt(office_ws):
    _make_excel(office_ws)
    result = await excel_query(
        {"path": "data.xlsx", "column": "score", "gt": 95, "limit": 3}, _ctx(office_ws),
    )
    assert not result.is_error
    assert "user96" in result.content
    assert "user0" not in result.content  # score 0 is not > 95


@pytest.mark.asyncio
async def test_excel_query_equals(office_ws):
    _make_excel(office_ws)
    result = await excel_query(
        {"path": "data.xlsx", "column": "city", "equals": "beijing", "limit": 2}, _ctx(office_ws),
    )
    assert not result.is_error
    assert "beijing" in result.content
    assert "shanghai" not in result.content.split("limited")[0]


@pytest.mark.asyncio
async def test_excel_query_column_not_found(office_ws):
    _make_excel(office_ws)
    result = await excel_query(
        {"path": "data.xlsx", "column": "nope", "equals": "x"}, _ctx(office_ws),
    )
    assert result.is_error
    assert "not found" in result.content


# ── Word ──


@pytest.mark.asyncio
async def test_word_extract_and_edit(office_ws):
    import docx

    d = docx.Document()
    d.add_paragraph("Hello World")
    d.add_paragraph("Old text here")
    d.save(str(office_ws / "doc.docx"))

    extracted = await word_extract({"path": "doc.docx"}, _ctx(office_ws))
    assert not extracted.is_error
    assert "Hello World" in extracted.content
    assert "Old text here" in extracted.content

    edited = await word_edit(
        {"path": "doc.docx", "find": "Old text", "replace": "New text"}, _ctx(office_ws),
    )
    assert not edited.is_error
    assert edited.metadata["replaced"] == 1

    re_extracted = await word_extract({"path": "doc.docx"}, _ctx(office_ws))
    assert "New text here" in re_extracted.content
    assert "Old text here" not in re_extracted.content


@pytest.mark.asyncio
async def test_word_edit_text_not_found(office_ws):
    import docx

    d = docx.Document()
    d.add_paragraph("Some text")
    d.save(str(office_ws / "doc.docx"))
    result = await word_edit(
        {"path": "doc.docx", "find": "missing", "replace": "x"}, _ctx(office_ws),
    )
    assert result.is_error
    assert "not found" in result.content


# ── PPT ──


@pytest.mark.asyncio
async def test_pptx_build_and_extract(office_ws):
    built = await pptx_build({"path": "slides.pptx", "slides": [
        {"title": "Slide One", "body": "First body"},
        {"title": "Slide Two", "body": "Second body"},
    ]}, _ctx(office_ws))
    assert not built.is_error
    assert built.metadata["slides"] == 2

    extracted = await pptx_extract({"path": "slides.pptx"}, _ctx(office_ws))
    assert not extracted.is_error
    assert "Slide One" in extracted.content
    assert "Slide Two" in extracted.content


@pytest.mark.asyncio
async def test_pptx_build_requires_slides(office_ws):
    result = await pptx_build({"path": "slides.pptx", "slides": []}, _ctx(office_ws))
    assert result.is_error
